import re
import csv
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===============================
# EXTRAIR NOMES + RM DO PDF
# ===============================
def extrair_nomes_de_pdf(caminho_pdf):
    leitor = PdfReader(caminho_pdf)
    texto = ""

    for pagina in leitor.pages:
        texto += pagina.extract_text() + "\n"

    linhas = texto.splitlines()
    alunos = []

    for linha in linhas:
        linha = linha.strip()

        # Nome em CAIXA ALTA
        if re.match(r'^[A-ZÁÉÍÓÚÃÕÂÊÎÔÛÇ ]{5,}$', linha):
            nome = ' '.join(linha.split())
            rm = ""

            # Procurar RM na mesma linha (6 a 8 números)
            rm_match = re.search(r'\b\d{6,8}\b', linha)
            if rm_match:
                rm = rm_match.group()

            alunos.append((nome, rm))

    print(f"✅ {len(alunos)} alunos extraídos do PDF.")
    return alunos


# ===============================
# GERAR CSV
# ===============================
def gerar_csv(nomes, turma, unidade, dominio="seudominio.com.br"):
    caminho_csv = f"usuarios_{turma}.csv"

    with open(caminho_csv, mode='w', newline='', encoding='utf-8') as arquivo:
        writer = csv.writer(arquivo, delimiter=';')

        writer.writerow(["Nome", "Usuario", "Email", "Turma", "Unidade"])

        for nome, rm in nomes:
            primeiro_nome = nome.split()[0].lower()
            usuario = primeiro_nome
            email = f"{usuario}@{dominio}"

            writer.writerow([nome, usuario, email, turma, unidade])

    print(f"📄 CSV gerado: {caminho_csv}")
    return caminho_csv


# ===============================
# GERAR WORD
# ===============================
def gerar_word_do_csv(caminho_csv, caminho_word, caminho_imagem, nomes_com_rm):
    doc = Document()

    for nome, rm in nomes_com_rm:
        doc.add_paragraph(f"Aluno: {nome}")
        doc.add_paragraph(f"RM: {rm}")
        doc.add_paragraph("Usuário: primeiro nome")
        doc.add_paragraph("Senha padrão: 123456")
        doc.add_paragraph("")

        try:
            doc.add_picture(caminho_imagem, width=Inches(2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

        doc.add_page_break()

    doc.save(caminho_word)
    print(f"📄 Word gerado: {caminho_word}")